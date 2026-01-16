"""
================================================================================
AASHTO 1993 Flexible Pavement Design - Streamlit Application
================================================================================
แอปพลิเคชันสำหรับออกแบบ Flexible Pavement ตามวิธี AASHTO 1993
เหมาะสำหรับการเรียนการสอนและการออกแบบจริง

Author: Civil Engineering Department
Version: 1.0
License: MIT
================================================================================
"""

import streamlit as st
import numpy as np
from scipy.optimize import brentq
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import BytesIO
import base64
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ================================================================================
# CONFIGURATION & CONSTANTS
# ================================================================================

st.set_page_config(
    page_title="AASHTO 1993 Flexible Pavement Design",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================================================================
# MATERIAL DATABASE
# ================================================================================

MATERIALS = {
    "Asphalt Concrete (AC)": {
        "elastic_modulus_mpa": 3100,      # MPa
        "elastic_modulus_psi": 450000,    # psi
        "layer_coeff": 0.44,              # a₁
        "drainage_range": (1.0, 1.0),     # m₁ typically 1.0 for AC
        "color": "#2C3E50",               # สีเข้มสำหรับแสดงผล
        "name_th": "แอสฟัลต์คอนกรีต (AC)"
    },
    "Crushed Aggregate Base": {
        "elastic_modulus_mpa": 207,       # MPa
        "elastic_modulus_psi": 30000,     # psi
        "layer_coeff": 0.14,              # a₂
        "drainage_range": (0.80, 1.25),   # m₂
        "color": "#795548",               # สีน้ำตาล
        "name_th": "ชั้นพื้นทางหินคลุก (Crushed Aggregate Base)"
    },
    "Cement Treated Base (CTB)": {
        "elastic_modulus_mpa": 690,       # MPa
        "elastic_modulus_psi": 100000,    # psi
        "layer_coeff": 0.23,              # a₂
        "drainage_range": (0.80, 1.25),   # m₂
        "color": "#78909C",               # สีเทา
        "name_th": "ชั้นพื้นทางซีเมนต์ปรับปรุงคุณภาพ (CTB)"
    },
    "Granular Subbase": {
        "elastic_modulus_mpa": 103,       # MPa
        "elastic_modulus_psi": 15000,     # psi
        "layer_coeff": 0.11,              # a₃
        "drainage_range": (0.80, 1.25),   # m₃
        "color": "#A1887F",               # สีน้ำตาลอ่อน
        "name_th": "ชั้นรองพื้นทางวัสดุมวลรวม (Granular Subbase)"
    },
    "Improved Subgrade": {
        "elastic_modulus_mpa": 69,        # MPa
        "elastic_modulus_psi": 10000,     # psi
        "layer_coeff": 0.08,              # a₄
        "drainage_range": (0.80, 1.25),   # m₄
        "color": "#BCAAA4",               # สีเบจ
        "name_th": "ชั้นดินเดิมปรับปรุง (Improved Subgrade)"
    },
    "Lime Treated Subgrade": {
        "elastic_modulus_mpa": 138,       # MPa
        "elastic_modulus_psi": 20000,     # psi
        "layer_coeff": 0.10,              # a₄
        "drainage_range": (0.80, 1.25),   # m₄
        "color": "#D7CCC8",               # สีครีม
        "name_th": "ชั้นดินเดิมปรับปรุงด้วยปูนขาว"
    },
    "Emulsified Asphalt Mix (EAM)": {
        "elastic_modulus_mpa": 2070,      # MPa
        "elastic_modulus_psi": 300000,    # psi
        "layer_coeff": 0.30,              # a₂
        "drainage_range": (1.0, 1.0),     # m₂
        "color": "#424242",               # สีเทาเข้ม
        "name_th": "ส่วนผสมแอสฟัลต์อิมัลชัน (EAM)"
    }
}

# ================================================================================
# RELIABILITY TABLE: Zr VALUES
# ================================================================================

RELIABILITY_ZR = {
    50: -0.000,
    60: -0.253,
    70: -0.524,
    75: -0.674,
    80: -0.841,
    85: -1.037,
    90: -1.282,
    91: -1.340,
    92: -1.405,
    93: -1.476,
    94: -1.555,
    95: -1.645,
    96: -1.751,
    97: -1.881,
    98: -2.054,
    99: -2.327,
    99.9: -3.090
}

# ================================================================================
# CORE CALCULATION FUNCTIONS
# ================================================================================

def aashto_1993_equation(SN: float, W18: float, Zr: float, So: float, 
                          delta_psi: float, Mr: float) -> float:
    """
    AASHTO 1993 Main Design Equation for Flexible Pavement
    
    สมการออกแบบหลักของ AASHTO 1993 สำหรับ Flexible Pavement
    
    Equation:
    log₁₀(W₁₈) = Zr×So + 9.36×log₁₀(SN+1) - 0.20 
                 + log₁₀(ΔPSI/(4.2-1.5)) / (0.4 + 1094/(SN+1)^5.19)
                 + 2.32×log₁₀(Mr) - 8.07
    
    Parameters:
    -----------
    SN : float
        Structural Number (trial value)
    W18 : float
        Design ESALs (18-kip equivalent single axle loads)
    Zr : float
        Standard normal deviate for reliability
    So : float
        Overall standard deviation
    delta_psi : float
        Change in serviceability (P₀ - Pₜ)
    Mr : float
        Subgrade resilient modulus (psi)
    
    Returns:
    --------
    float
        Residual of equation (should be zero when SN is correct)
    """
    # Left side of equation
    log_W18 = np.log10(W18)
    
    # Right side terms
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    
    # Serviceability loss term
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    
    # Subgrade modulus term
    term4 = 2.32 * np.log10(Mr) - 8.07
    
    # Right side total
    right_side = term1 + term2 + term3 + term4
    
    # Return residual (should be zero at correct SN)
    return right_side - log_W18


def calculate_sn_required(W18: float, Zr: float, So: float, 
                           delta_psi: float, Mr: float) -> dict:
    """
    Calculate required Structural Number (SN) using iterative solution
    
    คำนวณค่า SN ที่ต้องการโดยใช้วิธีการหาคำตอบแบบ iterative
    
    Uses Brent's method to solve AASHTO 1993 equation for SN
    
    Parameters:
    -----------
    W18 : float
        Design ESALs
    Zr : float
        Standard normal deviate
    So : float
        Overall standard deviation
    delta_psi : float
        Change in serviceability
    Mr : float
        Subgrade resilient modulus (psi)
    
    Returns:
    --------
    dict
        Contains SN_required and iteration details
    """
    # Define function to solve
    def f(SN):
        return aashto_1993_equation(SN, W18, Zr, So, delta_psi, Mr)
    
    # Find SN using Brent's method
    # Search range: 0.5 to 20 (typical practical range)
    try:
        SN_required = brentq(f, 0.1, 20.0, xtol=1e-6, maxiter=100)
        
        # Store iteration details
        iterations = []
        for SN_test in np.linspace(0.5, SN_required * 1.5, 10):
            residual = f(SN_test)
            iterations.append({
                'SN_test': SN_test,
                'residual': residual
            })
        
        return {
            'SN_required': round(SN_required, 3),
            'converged': True,
            'iterations': iterations
        }
    
    except ValueError as e:
        # If solution not found in range
        return {
            'SN_required': None,
            'converged': False,
            'error': str(e),
            'iterations': []
        }


def calculate_sn_provided(layers: list) -> dict:
    """
    Calculate provided Structural Number from layer configuration
    
    คำนวณค่า SN ที่ได้จากชั้นโครงสร้างถนนที่เลือก
    
    Equation: SN_provided = Σ(aᵢ × Dᵢ × mᵢ)
    
    Parameters:
    -----------
    layers : list of dict
        Each dict contains:
        - material: str (material name)
        - thickness_cm: float (thickness in cm)
        - drainage_coeff: float (mᵢ)
    
    Returns:
    --------
    dict
        Contains SN_provided and layer-by-layer breakdown
    """
    layer_details = []
    total_sn = 0.0
    
    for i, layer in enumerate(layers):
        material = layer['material']
        thickness_cm = layer['thickness_cm']
        drainage_coeff = layer['drainage_coeff']
        
        # Get material properties
        mat_props = MATERIALS[material]
        a_i = mat_props['layer_coeff']
        
        # Convert thickness: cm → inches
        thickness_inch = thickness_cm / 2.54
        
        # Calculate layer contribution to SN
        # For surface layer (AC), m = 1.0 typically
        if i == 0 and material == "Asphalt Concrete (AC)":
            m_i = 1.0  # Surface layer doesn't use drainage coefficient
        else:
            m_i = drainage_coeff
        
        sn_contribution = a_i * thickness_inch * m_i
        total_sn += sn_contribution
        
        layer_details.append({
            'layer_no': i + 1,
            'material': material,
            'material_th': mat_props['name_th'],
            'thickness_cm': thickness_cm,
            'thickness_inch': round(thickness_inch, 3),
            'layer_coeff': a_i,
            'drainage_coeff': m_i,
            'sn_contribution': round(sn_contribution, 4)
        })
    
    return {
        'SN_provided': round(total_sn, 3),
        'layer_details': layer_details
    }


def check_design(sn_required: float, sn_provided: float) -> dict:
    """
    Check if design is adequate
    
    ตรวจสอบว่าการออกแบบเพียงพอหรือไม่
    
    Parameters:
    -----------
    sn_required : float
        Required Structural Number from AASHTO equation
    sn_provided : float
        Provided Structural Number from layer configuration
    
    Returns:
    --------
    dict
        Pass/Fail status and safety margin
    """
    if sn_required is None:
        return {
            'status': 'ERROR',
            'message': 'Cannot calculate SN_required',
            'safety_margin': None
        }
    
    safety_margin = sn_provided - sn_required
    passed = sn_provided >= sn_required
    
    return {
        'status': 'PASS ✅' if passed else 'FAIL ❌',
        'passed': passed,
        'safety_margin': round(safety_margin, 3),
        'message': f"SN_provided ({sn_provided:.3f}) {'≥' if passed else '<'} SN_required ({sn_required:.3f})"
    }


# ================================================================================
# VISUALIZATION FUNCTIONS
# ================================================================================

def plot_pavement_section(layers: list, fig_width: float = 10, fig_height: float = 10, 
                          use_thai: bool = False) -> plt.Figure:
    """
    Draw vertical pavement section diagram
    
    วาดภาพหน้าตัดโครงสร้างถนนแนวตั้ง
    
    Parameters:
    -----------
    layers : list of dict
        Layer configuration with materials and thicknesses
    fig_width, fig_height : float
        Figure dimensions in inches
    use_thai : bool
        Use Thai text (requires Thai font installed)
    
    Returns:
    --------
    matplotlib.Figure
        Pavement section diagram
    """
    # Short names for display
    SHORT_NAMES = {
        "Asphalt Concrete (AC)": "AC",
        "Crushed Aggregate Base": "CAB",
        "Cement Treated Base (CTB)": "CTB",
        "Granular Subbase": "GSB",
        "Improved Subgrade": "ISG",
        "Lime Treated Subgrade": "LTS",
        "Emulsified Asphalt Mix (EAM)": "EAM"
    }
    
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    
    # Calculate total thickness for scaling
    total_thickness = sum([layer['thickness_cm'] for layer in layers])
    
    # Add subgrade thickness for visualization
    subgrade_thickness = total_thickness * 0.3
    
    # Drawing parameters
    layer_width = 6
    x_center = 4
    current_y = total_thickness + subgrade_thickness
    
    # Draw each layer from top to bottom
    for i, layer in enumerate(layers):
        material = layer['material']
        thickness_cm = layer['thickness_cm']
        mat_props = MATERIALS[material]
        
        # Create rectangle for layer
        rect = mpatches.FancyBboxPatch(
            (x_center - layer_width/2, current_y - thickness_cm),
            layer_width,
            thickness_cm,
            boxstyle="round,pad=0.02",
            facecolor=mat_props['color'],
            edgecolor='black',
            linewidth=2
        )
        ax.add_patch(rect)
        
        # Add layer information text (use short name for better display)
        short_name = SHORT_NAMES.get(material, material[:10])
        ax.text(x_center, current_y - thickness_cm/2, 
                short_name,
                ha='center', va='center',
                fontsize=12, fontweight='bold',
                color='white' if mat_props['color'] in ['#2C3E50', '#424242', '#78909C', '#795548'] else 'black')
        
        # Thickness label on the right
        ax.annotate(
            f'{thickness_cm:.0f} cm\n({thickness_cm/2.54:.1f}")',
            xy=(x_center + layer_width/2 + 0.1, current_y - thickness_cm/2),
            xytext=(x_center + layer_width/2 + 2.0, current_y - thickness_cm/2),
            fontsize=10, fontweight='bold',
            arrowprops=dict(arrowstyle='->', color='black', lw=1.5),
            va='center', ha='left',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='lightyellow', edgecolor='gray')
        )
        
        # Layer coefficient on the left
        ax.text(x_center - layer_width/2 - 0.5, current_y - thickness_cm/2,
                f'a{i+1}={mat_props["layer_coeff"]:.2f}',
                ha='right', va='center', fontsize=11, fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='lightblue', edgecolor='gray'))
        
        current_y -= thickness_cm
    
    # Draw subgrade
    subgrade_rect = mpatches.FancyBboxPatch(
        (x_center - layer_width/2, current_y - subgrade_thickness),
        layer_width,
        subgrade_thickness,
        boxstyle="round,pad=0.02",
        facecolor='#8D6E63',
        edgecolor='black',
        linewidth=2,
        linestyle='--'
    )
    ax.add_patch(subgrade_rect)
    ax.text(x_center, current_y - subgrade_thickness/2,
            'Subgrade',
            ha='center', va='center',
            fontsize=10, fontweight='bold', color='white')
    
    # Add title
    ax.text(x_center, total_thickness + subgrade_thickness + 2,
            'Pavement Cross Section',
            ha='center', va='center',
            fontsize=14, fontweight='bold')
    
    # Configure axes
    ax.set_xlim(-2, 14)
    ax.set_ylim(-5, total_thickness + subgrade_thickness + 4)
    ax.set_aspect('equal')
    ax.axis('off')
    
    plt.tight_layout()
    return fig


def get_figure_as_bytes(fig: plt.Figure) -> BytesIO:
    """
    Convert matplotlib figure to bytes for export
    
    แปลงรูปภาพ matplotlib เป็น bytes สำหรับการ export
    """
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    return buf


# ================================================================================
# WORD EXPORT FUNCTION
# ================================================================================

def create_word_report(project_title: str, inputs: dict, layers: list,
                       sn_required: float, sn_provided: dict,
                       design_check: dict, fig: plt.Figure) -> BytesIO:
    """
    Create Word document report for pavement design
    
    สร้างรายงาน Word สำหรับการออกแบบโครงสร้างถนน
    
    Parameters:
    -----------
    project_title : str
        Project title
    inputs : dict
        Design input parameters
    layers : list
        Layer configuration
    sn_required : float
        Required Structural Number
    sn_provided : dict
        Provided SN details
    design_check : dict
        Design check results
    fig : plt.Figure
        Pavement section figure
    
    Returns:
    --------
    BytesIO
        Word document as bytes
    """
    doc = Document()
    
    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    
    # Title
    title = doc.add_heading('รายงานการออกแบบ Flexible Pavement', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project Title
    doc.add_heading(f'โครงการ: {project_title}', level=1)
    
    # Date
    doc.add_paragraph(f'วันที่ออกแบบ: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    # Section 1: Design Method
    doc.add_heading('1. วิธีการออกแบบ', level=2)
    doc.add_paragraph(
        'การออกแบบโครงสร้างถนนใช้วิธี AASHTO 1993 Guide for Design of Pavement Structures '
        'โดยใช้สมการหลักดังนี้:'
    )
    
    # AASHTO Equation (simplified text version)
    eq_para = doc.add_paragraph()
    eq_para.add_run(
        'log₁₀(W₁₈) = Zᵣ×Sₒ + 9.36×log₁₀(SN+1) - 0.20 + '
        'log₁₀(ΔPSI/(4.2-1.5))/(0.4+1094/(SN+1)⁵·¹⁹) + 2.32×log₁₀(Mᵣ) - 8.07'
    ).italic = True
    
    # Section 2: Input Parameters
    doc.add_heading('2. ข้อมูลนำเข้าสำหรับการออกแบบ (Design Inputs)', level=2)
    
    # Create input table
    input_table = doc.add_table(rows=1, cols=3)
    input_table.style = 'Table Grid'
    
    # Header row
    header_cells = input_table.rows[0].cells
    header_cells[0].text = 'พารามิเตอร์'
    header_cells[1].text = 'ค่า'
    header_cells[2].text = 'หน่วย'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add input data
    input_data = [
        ('Design ESALs (W₁₈)', f'{inputs["W18"]:,.0f}', '18-kip ESAL'),
        ('Reliability (R)', f'{inputs["reliability"]}', '%'),
        ('Standard Normal Deviate (Zᵣ)', f'{inputs["Zr"]:.3f}', '-'),
        ('Overall Standard Deviation (Sₒ)', f'{inputs["So"]:.2f}', '-'),
        ('Initial Serviceability (P₀)', f'{inputs["P0"]:.1f}', '-'),
        ('Terminal Serviceability (Pₜ)', f'{inputs["Pt"]:.1f}', '-'),
        ('ΔPSI', f'{inputs["delta_psi"]:.1f}', '-'),
        ('Subgrade Resilient Modulus (Mᵣ)', f'{inputs["Mr"]:,.0f}', 'psi'),
    ]
    
    for param, value, unit in input_data:
        row = input_table.add_row()
        row.cells[0].text = param
        row.cells[1].text = value
        row.cells[2].text = unit
    
    # Section 3: Layer Configuration
    doc.add_heading('3. โครงสร้างชั้นทาง (Pavement Layer Configuration)', level=2)
    
    layer_table = doc.add_table(rows=1, cols=6)
    layer_table.style = 'Table Grid'
    
    # Header
    headers = ['ชั้นที่', 'วัสดุ', 'ความหนา (cm)', 'ความหนา (in)', 'aᵢ', 'mᵢ']
    header_cells = layer_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add layer data
    for detail in sn_provided['layer_details']:
        row = layer_table.add_row()
        row.cells[0].text = str(detail['layer_no'])
        row.cells[1].text = detail['material_th']
        row.cells[2].text = f'{detail["thickness_cm"]:.1f}'
        row.cells[3].text = f'{detail["thickness_inch"]:.3f}'
        row.cells[4].text = f'{detail["layer_coeff"]:.2f}'
        row.cells[5].text = f'{detail["drainage_coeff"]:.2f}'
    
    # Section 4: SN Calculation
    doc.add_heading('4. การคำนวณ Structural Number', level=2)
    
    doc.add_heading('4.1 SN Required (ค่าที่ต้องการ)', level=3)
    doc.add_paragraph(f'จากการแก้สมการ AASHTO 1993 โดยวิธี iterative:')
    doc.add_paragraph(f'SN_required = {sn_required:.3f}')
    
    doc.add_heading('4.2 SN Provided (ค่าที่ได้จากชั้นทาง)', level=3)
    doc.add_paragraph('สูตร: SN_provided = Σ(aᵢ × Dᵢ × mᵢ)')
    
    # SN contribution table
    sn_table = doc.add_table(rows=1, cols=5)
    sn_table.style = 'Table Grid'
    
    headers = ['ชั้นที่', 'aᵢ', 'Dᵢ (in)', 'mᵢ', 'SN contribution']
    header_cells = sn_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for detail in sn_provided['layer_details']:
        row = sn_table.add_row()
        row.cells[0].text = str(detail['layer_no'])
        row.cells[1].text = f'{detail["layer_coeff"]:.2f}'
        row.cells[2].text = f'{detail["thickness_inch"]:.3f}'
        row.cells[3].text = f'{detail["drainage_coeff"]:.2f}'
        row.cells[4].text = f'{detail["sn_contribution"]:.4f}'
    
    # Total row
    total_row = sn_table.add_row()
    total_row.cells[0].text = 'รวม'
    total_row.cells[4].text = f'{sn_provided["SN_provided"]:.3f}'
    for cell in total_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Section 5: Design Verification
    doc.add_heading('5. ผลการตรวจสอบการออกแบบ (Design Verification)', level=2)
    
    result_table = doc.add_table(rows=4, cols=2)
    result_table.style = 'Table Grid'
    
    result_data = [
        ('SN Required', f'{sn_required:.3f}'),
        ('SN Provided', f'{sn_provided["SN_provided"]:.3f}'),
        ('Safety Margin', f'{design_check["safety_margin"]:.3f}'),
        ('ผลการตรวจสอบ', design_check['status']),
    ]
    
    for i, (param, value) in enumerate(result_data):
        result_table.rows[i].cells[0].text = param
        result_table.rows[i].cells[1].text = value
    
    # Section 6: Pavement Section Figure
    doc.add_heading('6. ภาพตัดขวางโครงสร้างถนน', level=2)
    
    # Save figure to bytes and add to document
    fig_bytes = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes, width=Inches(5))
    
    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 7: Conclusion
    doc.add_heading('7. สรุปผลการออกแบบ', level=2)
    
    if design_check['passed']:
        conclusion = (
            f'การออกแบบ Flexible Pavement ผ่านเกณฑ์ที่กำหนด โดย SN_provided ({sn_provided["SN_provided"]:.3f}) '
            f'มากกว่าหรือเท่ากับ SN_required ({sn_required:.3f}) '
            f'ค่า Safety Margin = {design_check["safety_margin"]:.3f}'
        )
    else:
        conclusion = (
            f'การออกแบบ Flexible Pavement ไม่ผ่านเกณฑ์ที่กำหนด โดย SN_provided ({sn_provided["SN_provided"]:.3f}) '
            f'น้อยกว่า SN_required ({sn_required:.3f}) '
            f'กรุณาปรับเพิ่มความหนาชั้นทางหรือเปลี่ยนวัสดุที่มี Layer Coefficient สูงกว่า'
        )
    
    doc.add_paragraph(conclusion)
    
    # Save document to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes


# ================================================================================
# STREAMLIT USER INTERFACE
# ================================================================================

def main():
    """Main Streamlit application"""
    
    # Header
    st.title("🛣️ AASHTO 1993 Flexible Pavement Design")
    st.markdown("""
    **แอปพลิเคชันออกแบบโครงสร้างทางแบบยืดหยุ่น (Flexible Pavement) ตามมาตรฐาน AASHTO 1993**
    
    สำหรับการเรียนการสอนและการออกแบบจริง | For Teaching and Professional Design
    """)
    
    st.markdown("---")
    
    # Sidebar: Project Information
    with st.sidebar:
        st.header("📋 ข้อมูลโครงการ")
        project_title = st.text_input(
            "ชื่อโครงการ",
            value="Flexible Pavement Design Project"
        )
        
        st.markdown("---")
        st.header("📚 Material Database")
        
        with st.expander("ดูข้อมูลวัสดุทั้งหมด"):
            for mat_name, props in MATERIALS.items():
                st.markdown(f"**{props['name_th']}**")
                st.markdown(f"- E = {props['elastic_modulus_mpa']:,} MPa ({props['elastic_modulus_psi']:,} psi)")
                st.markdown(f"- a = {props['layer_coeff']}")
                st.markdown(f"- m range: {props['drainage_range']}")
                st.markdown("---")
    
    # Main content: Two columns
    col1, col2 = st.columns([1, 1])
    
    # ========================================
    # COLUMN 1: Design Inputs
    # ========================================
    with col1:
        st.header("📝 Design Inputs")
        
        # Traffic & Reliability Section
        st.subheader("1️⃣ Traffic & Reliability")
        
        W18 = st.number_input(
            "Design ESALs (W₁₈)",
            min_value=1000,
            max_value=100000000,
            value=5000000,
            step=100000,
            format="%d",
            help="จำนวน 18-kip ESAL ตลอดอายุการใช้งาน"
        )
        
        reliability = st.selectbox(
            "Reliability Level (R)",
            options=list(RELIABILITY_ZR.keys()),
            index=list(RELIABILITY_ZR.keys()).index(90),
            help="ระดับความเชื่อมั่นในการออกแบบ"
        )
        Zr = RELIABILITY_ZR[reliability]
        st.info(f"Zᵣ = {Zr:.3f}")
        
        So = st.number_input(
            "Overall Standard Deviation (Sₒ)",
            min_value=0.30,
            max_value=0.60,
            value=0.45,
            step=0.01,
            format="%.2f",
            help="ค่าเบี่ยงเบนมาตรฐานรวม (typical: 0.40-0.50 for flexible)"
        )
        
        # Serviceability Section
        st.subheader("2️⃣ Serviceability")
        
        col1a, col1b = st.columns(2)
        with col1a:
            P0 = st.number_input(
                "Initial Serviceability (P₀)",
                min_value=3.0,
                max_value=5.0,
                value=4.2,
                step=0.1,
                format="%.1f",
                help="ค่า Serviceability เริ่มต้น (typical: 4.2)"
            )
        
        with col1b:
            Pt = st.number_input(
                "Terminal Serviceability (Pₜ)",
                min_value=1.5,
                max_value=3.5,
                value=2.5,
                step=0.1,
                format="%.1f",
                help="ค่า Serviceability สิ้นสุด (typical: 2.0-2.5)"
            )
        
        delta_psi = P0 - Pt
        st.success(f"ΔPSI = P₀ - Pₜ = {P0:.1f} - {Pt:.1f} = **{delta_psi:.1f}**")
        
        # Subgrade Section
        st.subheader("3️⃣ Subgrade")
        
        Mr_input_type = st.radio(
            "วิธีกรอกค่า Subgrade Resilient Modulus",
            options=["กรอกค่า Mr โดยตรง (psi)", "คำนวณจาก CBR"],
            horizontal=True
        )
        
        if Mr_input_type == "กรอกค่า Mr โดยตรง (psi)":
            Mr = st.number_input(
                "Subgrade Resilient Modulus (Mᵣ) [psi]",
                min_value=1000,
                max_value=50000,
                value=5000,
                step=500,
                format="%d",
                help="ค่า Resilient Modulus ของดินเดิม"
            )
        else:
            CBR = st.number_input(
                "CBR (%)",
                min_value=1.0,
                max_value=30.0,
                value=5.0,
                step=0.5,
                format="%.1f"
            )
            # AASHTO correlation: Mr = 1500 × CBR (for CBR ≤ 10%)
            # or Mr = 2555 × CBR^0.64 (more general)
            if CBR <= 10:
                Mr = int(1500 * CBR)
            else:
                Mr = int(2555 * (CBR ** 0.64))
            st.info(f"Mᵣ (คำนวณจาก CBR) = **{Mr:,} psi** ({Mr * 0.006895:.1f} MPa)")
    
    # ========================================
    # COLUMN 2: Layer Configuration
    # ========================================
    with col2:
        st.header("🏗️ Layer Configuration")
        st.subheader("4️⃣ Pavement Layers")
        
        num_layers = st.slider(
            "จำนวนชั้นทาง",
            min_value=2,
            max_value=5,
            value=3,
            help="เลือกจำนวนชั้นทาง (2-5 ชั้น)"
        )
        
        layers = []
        
        for i in range(num_layers):
            st.markdown(f"##### ชั้นที่ {i+1}")
            
            col2a, col2b, col2c = st.columns([2, 1, 1])
            
            with col2a:
                material = st.selectbox(
                    f"วัสดุ",
                    options=list(MATERIALS.keys()),
                    index=min(i, len(MATERIALS)-1),
                    key=f"material_{i}",
                    help="เลือกประเภทวัสดุ"
                )
            
            with col2b:
                thickness = st.number_input(
                    f"ความหนา (cm)",
                    min_value=1.0,
                    max_value=50.0,
                    value=10.0 if i == 0 else (20.0 if i == 1 else 15.0),
                    step=1.0,
                    key=f"thickness_{i}",
                    format="%.1f"
                )
            
            with col2c:
                mat_props = MATERIALS[material]
                m_range = mat_props['drainage_range']
                
                drainage = st.number_input(
                    f"mᵢ",
                    min_value=m_range[0],
                    max_value=m_range[1],
                    value=(m_range[0] + m_range[1]) / 2 if i > 0 else 1.0,
                    step=0.05,
                    key=f"drainage_{i}",
                    format="%.2f",
                    help=f"Drainage coefficient (range: {m_range[0]:.2f} - {m_range[1]:.2f})"
                )
            
            # Display material properties
            st.caption(
                f"E = {mat_props['elastic_modulus_mpa']:,} MPa ({mat_props['elastic_modulus_psi']:,} psi) | "
                f"a = {mat_props['layer_coeff']}"
            )
            
            layers.append({
                'material': material,
                'thickness_cm': thickness,
                'drainage_coeff': drainage
            })
            
            st.markdown("---")
    
    # ========================================
    # CALCULATION & RESULTS
    # ========================================
    st.header("📊 Results & Verification")
    
    # Store inputs for report
    inputs = {
        'W18': W18,
        'reliability': reliability,
        'Zr': Zr,
        'So': So,
        'P0': P0,
        'Pt': Pt,
        'delta_psi': delta_psi,
        'Mr': Mr
    }
    
    # Calculate SN required
    sn_req_result = calculate_sn_required(W18, Zr, So, delta_psi, Mr)
    sn_required = sn_req_result['SN_required']
    
    # Calculate SN provided
    sn_prov_result = calculate_sn_provided(layers)
    sn_provided = sn_prov_result['SN_provided']
    
    # Check design
    design_check = check_design(sn_required, sn_provided)
    
    # Display results in columns
    res_col1, res_col2, res_col3 = st.columns(3)
    
    with res_col1:
        st.metric(
            label="SN Required",
            value=f"{sn_required:.3f}" if sn_required else "Error"
        )
    
    with res_col2:
        st.metric(
            label="SN Provided",
            value=f"{sn_provided:.3f}"
        )
    
    with res_col3:
        if design_check['passed']:
            st.metric(
                label="Safety Margin",
                value=f"{design_check['safety_margin']:.3f}",
                delta="PASS"
            )
        else:
            st.metric(
                label="Safety Margin",
                value=f"{design_check['safety_margin']:.3f}",
                delta="FAIL",
                delta_color="inverse"
            )
    
    # Status message
    if design_check['passed']:
        st.success(f"✅ **PASS**: {design_check['message']}")
    else:
        st.error(f"❌ **FAIL**: {design_check['message']}")
    
    # ========================================
    # PAVEMENT SECTION VISUALIZATION
    # ========================================
    st.header("📐 Pavement Section Diagram")
    
    fig = plot_pavement_section(layers)
    st.pyplot(fig)
    
    # ========================================
    # STEP-BY-STEP CALCULATION (Expandable)
    # ========================================
    with st.expander("📖 แสดงขั้นตอนการคำนวณ (Step-by-Step Calculation)"):
        st.subheader("Step 1: Input Parameters")
        st.markdown(f"""
        | Parameter | Value | Unit |
        |-----------|-------|------|
        | W₁₈ | {W18:,} | 18-kip ESAL |
        | R | {reliability} | % |
        | Zᵣ | {Zr:.3f} | - |
        | Sₒ | {So:.2f} | - |
        | P₀ | {P0:.1f} | - |
        | Pₜ | {Pt:.1f} | - |
        | ΔPSI | {delta_psi:.1f} | - |
        | Mᵣ | {Mr:,} | psi |
        """)
        
        st.subheader("Step 2: AASHTO 1993 Equation")
        st.latex(r'''
        \log_{10}(W_{18}) = Z_r S_o + 9.36 \log_{10}(SN+1) - 0.20 
        + \frac{\log_{10}(\Delta PSI / (4.2-1.5))}{0.4 + \frac{1094}{(SN+1)^{5.19}}}
        + 2.32 \log_{10}(M_r) - 8.07
        ''')
        
        st.subheader("Step 3: Iterative Solution for SN_required")
        st.markdown(f"""
        Using Brent's method to solve for SN:
        - **Initial search range**: 0.1 to 20.0
        - **Tolerance**: 1×10⁻⁶
        - **Solution**: SN_required = **{sn_required:.3f}**
        """)
        
        st.subheader("Step 4: Layer SN Contributions")
        st.markdown("**Formula**: SN_provided = Σ(aᵢ × Dᵢ × mᵢ)")
        
        # Layer table
        layer_data = []
        for detail in sn_prov_result['layer_details']:
            layer_data.append({
                'Layer': detail['layer_no'],
                'Material': detail['material'],
                'D (cm)': detail['thickness_cm'],
                'D (in)': f"{detail['thickness_inch']:.3f}",
                'aᵢ': detail['layer_coeff'],
                'mᵢ': detail['drainage_coeff'],
                'SN contribution': f"{detail['sn_contribution']:.4f}"
            })
        
        st.table(layer_data)
        
        st.markdown(f"**Total SN_provided = {sn_provided:.3f}**")
        
        st.subheader("Step 5: Design Verification")
        st.markdown(f"""
        | Check | Value |
        |-------|-------|
        | SN_required | {sn_required:.3f} |
        | SN_provided | {sn_provided:.3f} |
        | SN_provided ≥ SN_required | {design_check['passed']} |
        | Safety Margin | {design_check['safety_margin']:.3f} |
        | **Result** | **{design_check['status']}** |
        """)
    
    # ========================================
    # EXPORT REPORT
    # ========================================
    st.header("📄 Export Report")
    
    col_export1, col_export2 = st.columns(2)
    
    with col_export1:
        if st.button("📝 Generate Word Report (.docx)", type="primary"):
            with st.spinner("กำลังสร้างรายงาน..."):
                doc_bytes = create_word_report(
                    project_title=project_title,
                    inputs=inputs,
                    layers=layers,
                    sn_required=sn_required,
                    sn_provided=sn_prov_result,
                    design_check=design_check,
                    fig=fig
                )
                
                st.download_button(
                    label="⬇️ Download Word Report",
                    data=doc_bytes,
                    file_name=f"AASHTO_Flexible_Design_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    with col_export2:
        # PNG download for figure
        fig_bytes = get_figure_as_bytes(fig)
        st.download_button(
            label="📸 Download Section Diagram (PNG)",
            data=fig_bytes,
            file_name=f"Pavement_Section_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            mime="image/png"
        )
    
    # ========================================
    # FOOTER
    # ========================================
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: gray;'>
    <p>AASHTO 1993 Flexible Pavement Design Application</p>
    <p>พัฒนาสำหรับการเรียนการสอนและการออกแบบจริง</p>
    <p>Based on: AASHTO Guide for Design of Pavement Structures (1993)</p>
    </div>
    """, unsafe_allow_html=True)


# ================================================================================
# ENTRY POINT
# ================================================================================

if __name__ == "__main__":
    main()
