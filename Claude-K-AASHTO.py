import streamlit as st
from PIL import Image, ImageDraw, ImageFont
import io
import os
import math
from datetime import datetime

# ============================================
# Helper Functions for Calibration
# ============================================

def interpolate_log_scale(pixel, p1, v1, p2, v2):
    """
    Interpolate value on logarithmic scale given two calibration points.
    pixel: current pixel position
    p1, v1: first calibration point (pixel, value)
    p2, v2: second calibration point (pixel, value)
    """
    if v1 <= 0 or v2 <= 0:
        return 0
    log_v1 = math.log10(v1)
    log_v2 = math.log10(v2)
    # Linear interpolation in log space
    t = (pixel - p1) / (p2 - p1) if p2 != p1 else 0
    log_v = log_v1 + t * (log_v2 - log_v1)
    return 10 ** log_v

def interpolate_linear_scale(pixel, p1, v1, p2, v2):
    """
    Interpolate value on linear scale given two calibration points.
    """
    t = (pixel - p1) / (p2 - p1) if p2 != p1 else 0
    return v1 + t * (v2 - v1)

# ============================================
# Word Report Generation
# ============================================

def generate_word_report(params, image_bytes):
    """Generate Word report with calculation steps."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None, "ไม่พบ library python-docx กรุณาติดตั้งด้วย: pip install python-docx"
    
    doc = Document()
    
    # Set Thai font for the document
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(14)
    
    # Title
    title = doc.add_heading('รายงานการคำนวณ Composite Modulus of Subgrade Reaction (k∞)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('ตามวิธี AASHTO 1993 Rigid Pavement Design')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f'วันที่: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Section 1: Input Parameters
    doc.add_heading('1. ค่าพารามิเตอร์ที่ใช้ในการคำนวณ', level=1)
    
    # Create input table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['พารามิเตอร์', 'ค่า', 'หน่วย']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ('Roadbed Soil Resilient Modulus (MR)', f"{params['MR']:,.0f}", 'psi'),
        ('Subbase Elastic Modulus (ESB)', f"{params['ESB']:,.0f}", 'psi'),
        ('Subbase Thickness (DSB)', f"{params['DSB']:.1f}", 'inches'),
        ('Composite Modulus of Subgrade Reaction (k∞)', f"{params['k_inf']:,.0f}", 'pci'),
    ]
    
    for i, (param, value, unit) in enumerate(data):
        table.rows[i+1].cells[0].text = param
        table.rows[i+1].cells[1].text = value
        table.rows[i+1].cells[2].text = unit
    
    doc.add_paragraph()
    
    # Section 2: Calculation Steps
    doc.add_heading('2. ขั้นตอนการหาค่า k∞ จาก Nomograph', level=1)
    
    steps = [
        f"ขั้นตอนที่ 1: เริ่มจากแกน Roadbed Soil Resilient Modulus (MR) = {params['MR']:,.0f} psi ลากเส้นแนวตั้งขึ้น",
        f"ขั้นตอนที่ 2: จากค่า Subbase Elastic Modulus (ESB) = {params['ESB']:,.0f} psi หาจุดตัดกับเส้นโค้ง",
        f"ขั้นตอนที่ 3: ลากเส้นแนวนอนไปทางขวาจนตัดกับแกน Subbase Thickness (DSB) = {params['DSB']:.1f} inches",
        "ขั้นตอนที่ 4: จากจุดตัด ลากเส้นแนวนอนต่อไปจนตัดกับ Turning Line",
        f"ขั้นตอนที่ 5: จากจุดตัดบน Turning Line ลากเส้นแนวตั้งลงมาอ่านค่า k∞ = {params['k_inf']:,.0f} pci"
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    
    # Section 3: Nomograph Image
    doc.add_heading('3. Nomograph Chart พร้อมเส้นการอ่านค่า', level=1)
    
    # Add image
    if image_bytes:
        doc.add_picture(io.BytesIO(image_bytes), width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Section 4: Notes
    doc.add_heading('4. หมายเหตุ', level=1)
    notes = [
        "ค่า k∞ ที่ได้เป็นค่า Composite Modulus of Subgrade Reaction สำหรับ Semi-infinite Subgrade Depth",
        "การอ่านค่าจาก Nomograph มีความคลาดเคลื่อนโดยธรรมชาติ ควรพิจารณาใช้ค่าที่เหมาะสมกับสภาพจริง",
        "Reference: AASHTO Guide for Design of Pavement Structures 1993, Figure 3.3"
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer, None

# ============================================
# Main Application
# ============================================

def main():
    st.set_page_config(
        page_title="AASHTO k∞ Nomograph Calculator",
        page_icon="📊",
        layout="wide"
    )
    
    st.title("🛣️ เครื่องมือหาค่า k∞ จาก Nomograph")
    st.markdown("**Composite Modulus of Subgrade Reaction ตามวิธี AASHTO 1993**")
    
    # Create tabs
    tab1, tab2, tab3 = st.tabs(["📊 อ่านค่าจากกราฟ", "⚙️ ตั้งค่า Calibration", "📋 คู่มือการใช้งาน"])
    
    # Initialize session state for calibration
    if 'calibration' not in st.session_state:
        st.session_state.calibration = {
            # Default calibration values (will be adjusted based on actual image)
            'MR_p1': 100, 'MR_v1': 1000,
            'MR_p2': 500, 'MR_v2': 20000,
            'ESB_p1': 100, 'ESB_v1': 15000,
            'ESB_p2': 400, 'ESB_v2': 1000000,
            'DSB_p1': 200, 'DSB_v1': 18,
            'DSB_p2': 400, 'DSB_v2': 4,
            'k_p1': 300, 'k_v1': 50,
            'k_p2': 100, 'k_v2': 2000,
        }
    
    # Tab 2: Calibration Settings
    with tab2:
        st.header("⚙️ ตั้งค่า Calibration")
        st.info("ปรับค่า Calibration ให้ตรงกับตำแหน่งแกนบนภาพ Nomograph ของคุณ")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("แกน MR (Roadbed Soil Resilient Modulus)")
            st.session_state.calibration['MR_p1'] = st.number_input("MR Pixel 1", value=st.session_state.calibration['MR_p1'], key='mr_p1')
            st.session_state.calibration['MR_v1'] = st.number_input("MR Value 1 (psi)", value=st.session_state.calibration['MR_v1'], key='mr_v1')
            st.session_state.calibration['MR_p2'] = st.number_input("MR Pixel 2", value=st.session_state.calibration['MR_p2'], key='mr_p2')
            st.session_state.calibration['MR_v2'] = st.number_input("MR Value 2 (psi)", value=st.session_state.calibration['MR_v2'], key='mr_v2')
            
            st.subheader("แกน ESB (Subbase Elastic Modulus)")
            st.session_state.calibration['ESB_p1'] = st.number_input("ESB Pixel 1", value=st.session_state.calibration['ESB_p1'], key='esb_p1')
            st.session_state.calibration['ESB_v1'] = st.number_input("ESB Value 1 (psi)", value=st.session_state.calibration['ESB_v1'], key='esb_v1')
            st.session_state.calibration['ESB_p2'] = st.number_input("ESB Pixel 2", value=st.session_state.calibration['ESB_p2'], key='esb_p2')
            st.session_state.calibration['ESB_v2'] = st.number_input("ESB Value 2 (psi)", value=st.session_state.calibration['ESB_v2'], key='esb_v2')
        
        with col2:
            st.subheader("แกน DSB (Subbase Thickness)")
            st.session_state.calibration['DSB_p1'] = st.number_input("DSB Pixel 1", value=st.session_state.calibration['DSB_p1'], key='dsb_p1')
            st.session_state.calibration['DSB_v1'] = st.number_input("DSB Value 1 (inches)", value=st.session_state.calibration['DSB_v1'], key='dsb_v1')
            st.session_state.calibration['DSB_p2'] = st.number_input("DSB Pixel 2", value=st.session_state.calibration['DSB_p2'], key='dsb_p2')
            st.session_state.calibration['DSB_v2'] = st.number_input("DSB Value 2 (inches)", value=st.session_state.calibration['DSB_v2'], key='dsb_v2')
            
            st.subheader("แกน k∞ (Composite Modulus)")
            st.session_state.calibration['k_p1'] = st.number_input("k∞ Pixel 1", value=st.session_state.calibration['k_p1'], key='k_p1')
            st.session_state.calibration['k_v1'] = st.number_input("k∞ Value 1 (pci)", value=st.session_state.calibration['k_v1'], key='k_v1')
            st.session_state.calibration['k_p2'] = st.number_input("k∞ Pixel 2", value=st.session_state.calibration['k_p2'], key='k_p2')
            st.session_state.calibration['k_v2'] = st.number_input("k∞ Value 2 (pci)", value=st.session_state.calibration['k_v2'], key='k_v2')
    
    # Tab 3: Instructions
    with tab3:
        st.header("📋 คู่มือการใช้งาน")
        
        st.subheader("วิธีการใช้งาน")
        st.markdown("""
        1. **อัปโหลดภาพ Nomograph** - ใช้ภาพ Figure 3.3 จาก AASHTO 1993
        2. **ปรับเส้น Turning Line (สีเขียว)** - ให้ทับกับเส้น Turning Line บนกราฟพอดี
        3. **กำหนดจุดเริ่มต้น** - ปรับตำแหน่งแกน MR และ ESB
        4. **ปรับจุดตัดเส้นโค้ง** - ให้เส้นสีแดงตัดกับเส้นโค้ง DSB ที่ต้องการ
        5. **อ่านค่า k∞** - ค่าจะแสดงจากจุดตัดบนแกน k∞
        6. **Export รายงาน** - กดปุ่มดาวน์โหลด Word Report
        """)
        
        st.subheader("ความหมายของสี")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("🟢 **เส้นสีเขียว** - Turning Line (เส้นอ้างอิง)")
            st.markdown("🔴 **เส้นสีแดง** - เส้นจาก MR ขึ้นไปหาจุดตัด ESB")
        with col2:
            st.markdown("🔵 **เส้นสีน้ำเงิน** - เส้นจาก ESB ไปหาจุดตัด DSB")
            st.markdown("🟠 **เส้นสีส้ม** - เส้นไปหา Turning Line และลงสู่แกน k∞")
        
        st.subheader("Reference")
        st.markdown("""
        - AASHTO Guide for Design of Pavement Structures 1993
        - Figure 3.3: Chart for Estimating Composite Modulus of Subgrade Reaction
        """)
    
    # Tab 1: Main Calculator
    with tab1:
        # File uploader
        uploaded_file = st.file_uploader(
            "📁 อัปโหลดภาพ Nomograph (Figure 3.3 AASHTO 1993)",
            type=['png', 'jpg', 'jpeg'],
            help="รองรับไฟล์ PNG, JPG, JPEG"
        )
        
        if uploaded_file is not None:
            image = Image.open(uploaded_file).convert("RGB")
            width, height = image.size
            img_draw = image.copy()
            draw = ImageDraw.Draw(img_draw)
            
            # Layout: sidebar for controls, main area for image
            st.sidebar.header("🎯 ตั้งค่าเส้นบนกราฟ")
            
            # =========================================
            # Section 1: Turning Line (Green)
            # =========================================
            with st.sidebar.expander("1️⃣ เส้น Turning Line (สีเขียว)", expanded=True):
                st.caption("ปรับให้ทับกับเส้น Turning Line บนกราฟ")
                
                col1, col2 = st.columns(2)
                with col1:
                    green_x1 = st.slider("X เริ่มต้น", 0, width, 421, key="gx1")
                    green_y1 = st.slider("Y เริ่มต้น", 0, height, 346, key="gy1")
                with col2:
                    green_x2 = st.slider("X สิ้นสุด", 0, width, 691, key="gx2")
                    green_y2 = st.slider("Y สิ้นสุด", 0, height, 620, key="gy2")
            
            # Draw green turning line
            draw.line([(green_x1, green_y1), (green_x2, green_y2)], fill="green", width=6)
            
            # Calculate slope
            if green_x2 - green_x1 == 0:
                st.error("⚠️ เส้น Turning Line ต้องไม่เป็นแนวตั้ง")
                return
            slope_green = (green_y2 - green_y1) / (green_x2 - green_x1)
            
            # =========================================
            # Section 2: Input Parameters
            # =========================================
            with st.sidebar.expander("2️⃣ ค่าพารามิเตอร์ (เส้นแดง/น้ำเงิน)", expanded=True):
                st.caption("กำหนดตำแหน่งจุดเริ่มต้นและจุดตัด")
                
                # Starting point (MR axis)
                start_x = st.slider(
                    "ตำแหน่ง MR (แนวนอน)", 
                    0, width, int(width * 0.25),
                    help="ตำแหน่งบนแกน Roadbed Soil Resilient Modulus"
                )
                start_y_bottom = st.slider(
                    "จุดเริ่มต้น (แนวตั้ง)", 
                    0, height, int(height * 0.85),
                    help="จุดเริ่มต้นที่ขอบล่างของกราฟ"
                )
                
                # Intersection point with ESB curve (for blue line)
                stop_y_esb = st.slider(
                    "จุดตัดเส้นโค้ง ESB (แนวตั้ง)", 
                    0, height, int(height * 0.15),
                    help="ความสูงของจุดตัดกับเส้นโค้ง Subbase Elastic Modulus"
                )
                
                # Intersection point (ESB/DSB curves)
                stop_y_1 = st.slider(
                    "จุดตัดเส้นโค้ง DSB (แนวตั้ง)", 
                    0, height, int(height * 0.35),
                    help="ความสูงของจุดตัดกับเส้นโค้ง Subbase Thickness"
                )
            
            # =========================================
            # Section 3: Output Settings
            # =========================================
            with st.sidebar.expander("3️⃣ แกน k∞ (เส้นส้ม)", expanded=True):
                k_axis_y = st.slider(
                    "ตำแหน่งแกน k∞ (แนวตั้ง)", 
                    0, height, int(height * 0.15),
                    help="ความสูงของแกน k∞ สำหรับอ่านค่า"
                )
            
            # =========================================
            # Calculate constrained point on turning line
            # =========================================
            target_y = stop_y_1
            constrained_x = green_x1 + (target_y - green_y1) / slope_green
            constrained_x = int(constrained_x)
            
            # =========================================
            # Draw Lines
            # =========================================
            line_width = 4
            arrow_size = 12
            
            # Red line: vertical from MR axis up to ESB intersection
            draw.line([(start_x, start_y_bottom), (start_x, stop_y_esb)], fill="red", width=line_width)
            
            # Blue line: horizontal from ESB intersection to DSB curve (same starting point as red line top)
            # Calculate the x position where blue line meets DSB curve at stop_y_1
            blue_end_x = start_x + int((stop_y_1 - stop_y_esb) * 1.5)  # Approximate, user can adjust via stop_y_1
            draw.line([(start_x, stop_y_esb), (blue_end_x, stop_y_1)], fill="blue", width=line_width)
            
            # Arrow for blue line
            draw.polygon([
                (blue_end_x, stop_y_1),
                (blue_end_x - arrow_size, stop_y_1 - arrow_size//2),
                (blue_end_x - arrow_size, stop_y_1 + arrow_size//2)
            ], fill="blue")
            
            # Orange line: horizontal from DSB intersection to turning line
            draw.line([(blue_end_x, stop_y_1), (constrained_x, stop_y_1)], fill="orange", width=line_width)
            
            # Orange line: vertical down to k∞ axis
            draw.line([(constrained_x, stop_y_1), (constrained_x, k_axis_y)], fill="orange", width=line_width)
            
            # Draw intersection point
            radius = 8
            draw.ellipse([
                (constrained_x - radius, stop_y_1 - radius),
                (constrained_x + radius, stop_y_1 + radius)
            ], fill="black", outline="white", width=2)
            
            # Draw arrows
            # Arrow at start (MR)
            draw.polygon([
                (start_x, start_y_bottom),
                (start_x - arrow_size//2, start_y_bottom - arrow_size),
                (start_x + arrow_size//2, start_y_bottom - arrow_size)
            ], fill="red")
            
            # Arrow at end (k∞)
            draw.polygon([
                (constrained_x, k_axis_y),
                (constrained_x - arrow_size//2, k_axis_y + arrow_size),
                (constrained_x + arrow_size//2, k_axis_y + arrow_size)
            ], fill="orange")
            
            # =========================================
            # Display Image
            # =========================================
            st.image(img_draw, caption="Nomograph พร้อมเส้นการอ่านค่า", use_container_width=True)
            
            # =========================================
            # Results Section
            # =========================================
            st.markdown("---")
            st.subheader("📊 ผลลัพธ์")
            
            # Display pixel coordinates
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("จุดเริ่มต้น (MR)", f"({start_x}, {start_y_bottom})")
            with col2:
                st.metric("จุดตัด Turning Line", f"({constrained_x}, {stop_y_1})")
            with col3:
                st.metric("จุดอ่านค่า k∞", f"({constrained_x}, {k_axis_y})")
            
            st.info(f"📍 พิกัดจุดตัดบน Turning Line: **({constrained_x}, {stop_y_1})**")
            
            # =========================================
            # Manual Value Input for Report
            # =========================================
            st.markdown("---")
            st.subheader("📝 กรอกค่าสำหรับรายงาน")
            st.caption("กรอกค่าที่อ่านได้จากกราฟเพื่อใช้ในการสร้างรายงาน")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                MR_value = st.number_input(
                    "MR (psi)",
                    min_value=1000,
                    max_value=20000,
                    value=7000,
                    step=500,
                    help="Roadbed Soil Resilient Modulus"
                )
            
            with col2:
                ESB_value = st.number_input(
                    "ESB (psi)",
                    min_value=15000,
                    max_value=1000000,
                    value=50000,
                    step=5000,
                    help="Subbase Elastic Modulus"
                )
            
            with col3:
                DSB_value = st.number_input(
                    "DSB (inches)",
                    min_value=0.0,
                    max_value=18.0,
                    value=6.0,
                    step=0.5,
                    help="Subbase Thickness"
                )
            
            with col4:
                k_inf_value = st.number_input(
                    "k∞ (pci)",
                    min_value=50,
                    max_value=2000,
                    value=500,
                    step=50,
                    help="Composite Modulus of Subgrade Reaction"
                )
            
            # Display summary
            st.success(f"""
            **สรุปค่าที่ใช้ในการออกแบบ:**
            - MR = {MR_value:,} psi
            - ESB = {ESB_value:,} psi  
            - DSB = {DSB_value} inches
            - **k∞ = {k_inf_value:,} pci**
            """)
            
            # =========================================
            # Export Section
            # =========================================
            st.markdown("---")
            st.subheader("📥 ดาวน์โหลดรายงาน")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Export image
                img_buffer = io.BytesIO()
                img_draw.save(img_buffer, format='PNG')
                img_bytes = img_buffer.getvalue()
                
                st.download_button(
                    label="📷 ดาวน์โหลดภาพ (PNG)",
                    data=img_bytes,
                    file_name=f"nomograph_k_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
                    mime="image/png"
                )
            
            with col2:
                # Export Word report
                params = {
                    'MR': MR_value,
                    'ESB': ESB_value,
                    'DSB': DSB_value,
                    'k_inf': k_inf_value,
                    'constrained_x': constrained_x,
                    'stop_y_1': stop_y_1
                }
                
                if st.button("📄 สร้าง Word Report"):
                    with st.spinner("กำลังสร้างรายงาน..."):
                        report_buffer, error = generate_word_report(params, img_bytes)
                        
                        if error:
                            st.error(error)
                        else:
                            st.download_button(
                                label="📥 ดาวน์โหลด Word Report",
                                data=report_buffer,
                                file_name=f"k_infinity_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            st.success("✅ สร้างรายงานสำเร็จ!")
        
        else:
            # Show placeholder when no image uploaded
            st.info("👆 กรุณาอัปโหลดภาพ Nomograph เพื่อเริ่มต้นใช้งาน")
            
            # Show example
            with st.expander("📖 ตัวอย่างการใช้งาน"):
                st.markdown("""
                1. อัปโหลดภาพ **Figure 3.3** จากหนังสือ AASHTO 1993
                2. ปรับ **เส้นสีเขียว** ให้ทับกับ Turning Line
                3. ปรับ **เส้นสีแดง** ตามค่า MR ที่ต้องการ
                4. ปรับ **จุดตัด** ให้ตรงกับเส้นโค้ง DSB
                5. อ่านค่า **k∞** จากตำแหน่งเส้นสีส้มตัดแกนขวา
                6. กรอกค่าและดาวน์โหลดรายงาน
                """)

if __name__ == "__main__":
    main()
